from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Set document styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper function to add a heading
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

# Helper function to add a bullet list
def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# Contact Info and Header
doc.add_paragraph("Luca Lorusso", style='Title')
doc.add_paragraph("📍 Selden, NY 11784 | 📞 (631) 645-1324 | ✉️ luca.lorusso@icloud.com")
doc.add_paragraph("🔗 LinkedIn: https://www.linkedin.com/in/luca-lorusso-qa-engineer")

# Summary
add_heading("Professional Summary", level=2)
doc.add_paragraph(
    "Detail-oriented QA Automation Engineer with over 3 years of experience ensuring the quality, stability, and performance of web "
    "and mobile applications. Proficient in manual and automation testing, with hands-on expertise in tools like Selenium, Postman, "
    "Charles Proxy, and Chrome DevTools. Adept in agile environments, bug tracking with JIRA, and cross-functional collaboration. "
    "Known for reducing application defects by 25% and maintaining a 95%+ customer satisfaction rate. Fluent in English, Italian, and Spanish."
)

# Technical Skills
add_heading("Technical Skills", level=2)
add_bullets([
    "Testing Tools: Selenium WebDriver, Chrome DevTools, Postman, Charles Proxy, BrowserStack",
    "Platforms: Android Studio, Xcode, macOS, iOS, Windows, Linux, Android",
    "Languages: Python (basic automation), SQL",
    "Test Management: JIRA, TestRail",
    "Testing Types: Manual, Functional, Regression, Smoke, UI/UX, Compatibility",
    "Others: ADB, UNIX, REST API Testing"
])

# Professional Experience
add_heading("Professional Experience", level=2)

doc.add_paragraph("Software QA Engineer", style='Heading 3')
doc.add_paragraph("Invoice Maker – Selden, NY | June 2023 – Present")
add_bullets([
    "Designed and executed test plans and test cases for web and mobile applications.",
    "Conducted cross-browser compatibility testing across Chrome, Safari, Firefox, and Edge.",
    "Identified and logged defects using JIRA, collaborating with developers to expedite resolution.",
    "Performed black-box, regression, UI/UX, and smoke testing throughout SDLC.",
    "Used Postman and Charles Proxy for API testing and debugging."
])

doc.add_paragraph("Software QA Engineer", style='Heading 3')
doc.add_paragraph("Go Global World | March 2022 – June 2023")
add_bullets([
    "Validated app functionality across multiple devices and platforms ensuring consistent UX.",
    "Authored comprehensive manual test cases aligned with project requirements.",
    "Performed API validation, browser testing, and defect lifecycle management in JIRA.",
    "Participated in sprint planning meetings, contributing QA feedback to the product roadmap."
])

doc.add_paragraph("Sales Specialist (Part-Time)", style='Heading 3')
doc.add_paragraph("Apple Inc. – Lake Grove, NY | September 2023 – Present")
add_bullets([
    "Delivered customer-centric service and exceeded sales KPIs.",
    "Fourth-quarter achievement: 100 TMS, 40 Net Promoter Score, 97 Product Zone NPS.",
    "Contributed to a diverse and inclusive team culture by embracing Apple’s core values."
])

# Certifications
add_heading("Certifications", level=2)
add_bullets([
    "Software Testing & QA Bootcamp – Careerist (2024)",
    "Selenium WebDriver with Python & Frameworks – Ongoing (2023)",
    "SQL Essential Training – LinkedIn Learning (2022)"
])

# Education
add_heading("Education", level=2)
doc.add_paragraph("B.A. in Culinary Arts\nI.P.S.S.A.R.T, Castellana Grotte – Bari, Italy")

# Languages
add_heading("Languages", level=2)
doc.add_paragraph("English (Fluent), Italian (Native), Spanish (Professional Proficiency)")

# Save the document
docx_path = "/mnt/data/Luca_Lorusso_QA_Engineer_Resume_2025.docx"
pdf_path = "/mnt/data/Luca_Lorusso_QA_Engineer_Resume_2025.pdf"
doc.save(docx_path)

docx_path  # Output the DOCX path for download first

